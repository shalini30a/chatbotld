import streamlit as st
import cohere
from docx import Document
import tempfile
import pymupdf
import csv
from datetime import datetime
import os
import re
from collections import Counter
# from importlib-metadata import version  # For Python >= 3.8

# Set your Cohere API key
COHERE_API_KEY = "njZaNmhiiUXOcP0gaThBFXwW8g5AdZRqG9r3SPAF"  # Replace with your actual API key
cohere_client = cohere.Client(COHERE_API_KEY)

# Initialize session state
if "conversation" not in st.session_state:
    st.session_state.conversation = []
if "exit" not in st.session_state:
    st.session_state.exit = False
if "faq_text" not in st.session_state:
    st.session_state.faq_text = ""
if "rerun_trigger" not in st.session_state:
    st.session_state.rerun_trigger = 0

def preprocess_text(text):
    # Lowercase and remove special characters
    text = re.sub(r'[^\w\s]', '', text.lower())
    return text.split()

def create_document_index(text):
    # Create a Counter object for quick keyword matching
    words = preprocess_text(text)
    return Counter(words)

# Create a login page
def login():
    st.title("Login")
    email = st.text_input("Enter your Email ID:")
    if email:
        st.session_state.email = email
        st.success("Login successful!")
        return email
    return None

# Store queries and responses in a CSV file
def store_query(email, query, response):
    file_path = 'user_queries.csv'
    file_exists = os.path.exists(file_path)

    # Open the file in append mode
    with open(file_path, 'a', newline='', encoding='utf-8') as file:
        writer = csv.writer(file)

        # Write the header if the file is new or empty
        if not file_exists or os.stat(file_path).st_size == 0:
            writer.writerow(["Email", "Query", "Response", "Timestamp"])

        # Write the query and response data
        writer.writerow([email, query, response, datetime.now()])
# Function to extract text and links from DOCX
# def extract_text_and_links_from_docx(docx_file):
#     doc = Document(docx_file)
#     text = ''
#     links = []
#
#     for paragraph in doc.paragraphs:
#         for run in paragraph.runs:
#             text += run.text
#         if paragraph._element.xpath(".//w:hyperlink"):
#             for hyperlink in paragraph._element.xpath(".//w:hyperlink"):
#                 rId = hyperlink.get("{http://schemas.openxmlformats.org/officeDocument/2006/relationships}id")
#                 if rId:
#                     part = doc.part
#                     rel = part.rels.get(rId)
#                     if rel:
#                         links.append(rel.target_ref)
#     return text, links
#
# # Function to process the uploaded FAQ file
# def process_faq_file(uploaded_file):
#     if uploaded_file.name.endswith('.docx'):
#         with tempfile.NamedTemporaryFile(delete=False) as tmp_file:
#             tmp_file.write(uploaded_file.getbuffer())
#             tmp_file_path = tmp_file.name
#         with open(tmp_file_path, "rb") as temp_docx:
#             return extract_text_and_links_from_docx(temp_docx)
#     else:
#         return "Unsupported file format. Please upload a DOCX.", []
# Function to extract text and links from DOCX
# File uploader for multiple files


def extract_text_and_links_from_pdf(pdf_file):
    pdf_reader = pymupdf.open(stream=pdf_file.read(), filetype="pdf")
    text = ""
    links = []
    for page in pdf_reader:
        text += page.get_text()
        links += page.get_links()  # Extracts links from the page
    pdf_reader.close()

    # Format links for display
    formatted_links = [
        f"[{link['uri']}]({link['uri']})" for link in links if link.get('uri')
    ]
    return text, formatted_links

# Function to extract text and links from DOCX
def extract_text_and_links_from_docx(docx_file):
    doc = Document(docx_file)
    text = ''
    links = []

    # Define XML namespace map for Word documents
    nsmap = {
        'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'
    }

    # Access the XML of the document
    for paragraph in doc.paragraphs:
        # Extracting text from the paragraph
        for run in paragraph.runs:
            text += run.text

        # Check for hyperlinks in the paragraph using the correct namespace
        for hyperlink in paragraph._element.findall(".//w:hyperlink", namespaces=nsmap):
            # Extract the relationship ID of the hyperlink
            rId = hyperlink.get("{http://schemas.openxmlformats.org/officeDocument/2006/relationships}id")
            if rId:
                # Get the target of the hyperlink from the document part relationships
                part = doc.part
                rel = part.rels.get(rId)
                if rel and hasattr(rel, 'target_ref'):
                    links.append(rel.target_ref)

    return text, links


# Function to process the uploaded FAQ file
def process_faq_file(uploaded_file):
    if uploaded_file.name.endswith('.docx'):
        with tempfile.NamedTemporaryFile(delete=False) as tmp_file:
            tmp_file.write(uploaded_file.getbuffer())
            tmp_file_path = tmp_file.name
        with open(tmp_file_path, "rb") as temp_docx:
            return extract_text_and_links_from_docx(temp_docx)
    else:
        return "Unsupported file format. Please upload a DOCX.", []


# Test the code with a sample DOCX file
# uploaded_file = st.sidebar.file_uploader("Upload FAQ file (DOCX)", type=["docx"])
# faq_text, links = process_faq_file(uploaded_file)

# Now faq_text will contain the text and links will contain the hyperlinks found in the document

# Function to get an AI-generated answer using Cohere
def get_answer_from_cohere(query, faq_text):
    try:
        prompt = (
            f"You are an assistant specialized in L&D-related queries. "
            f"Here is an FAQ document:\n{faq_text}\n\n"
            f"Answer the following query:\n{query}"
        )
        response = cohere_client.generate(
            model='command-xlarge-nightly',
            prompt=prompt,
            max_tokens=300,
            temperature=0.5,
        )
        return response.generations[0].text.strip()
    except Exception as e:
        return f"Error: {e}"
def is_query_relevant(query, document_index):
    query_words = preprocess_text(query)
    match_count = sum(1 for word in query_words if word in document_index)
    return match_count > 0  # Consider the query relevant if it matches at least one word
# Streamlit UI
st.markdown("""
    <style>
    body {
        background-color: "#FFFFFF";
        color: #000000;
    }
    .response {
        color: #008000;  /* Blue color for answers */
        font-size: 18px;
    }
    .question {
        color: #1E90FF;  /* Green color for questions */
        font-size: 18px;
    }
    .exit-text {
        color: #FF6347;  /* Tomato color for exit message */
        font-size: 20px;
    }
    .wel-text {
        color: #800080;  /* Tomato color for exit message */
        font-size: 18px;
    }
    .title {
        color: #1E90FF;  /* Green color for questions */
        font-size: 35px;
    }
    </style>
    """, unsafe_allow_html=True)
col1, col2 = st.columns([2,1])  # Adjust the proportions as needed
with col1:
    st.markdown("# <span class='title'> Indi-GO-LanD   🤖 </span>",unsafe_allow_html=True)
st.markdown("##### <span class='wel-text'> Welocome to Your Learning Assistant !! Do you have queries with your Learning & Development Department? </span> ",unsafe_allow_html=True )

with col2:
    st.image("indiumlogo.png", width=400)
# Main Streamlit App
# email = None
# if "email" not in st.session_state:
#     email = login()
#
# if email or "email" in st.session_state:
#     email = email or st.session_state.email

# st.sidebar.title("Settings")
# # theme = st.sidebar.radio("Select Theme:", ["Light", "Dark"])
# # st.sidebar.radio("Select Theme:", ["Light", "Dark"])
# # Apply the chosen theme
# # theme = st.sidebar.radio("Select Theme:", ["Light", "Dark"])
# # if theme == "Light":
# #     st.markdown("<span class='question'> </span>", unsafe_allow_html=True)
# # else:
# #     st.markdown("<span class='answer'> </span>", unsafe_allow_html=True)
# # File uploader and document info display in sidebar
# # uploaded_file = st.sidebar.file_uploader("Upload FAQ file (DOCX)", type=["docx"])
#
# uploaded_files = st.sidebar.file_uploader("Upload multiple files", type=["docx", "pdf", "txt"], accept_multiple_files=True)
#
# # Check if any files are uploaded
# if uploaded_files:
#     for uploaded_file in uploaded_files:
#         faq_text, links = process_faq_file(uploaded_file)
#         # Display file name and content (for example, display the first 100 characters of the file content)
#         st.sidebar.write(f"File name: {uploaded_file.name}")
#         st.sidebar.write(f"Number of paragraphs in FAQ: {len(faq_text.splitlines())}")
#         # file_content = uploaded_file.read().decode("utf-8", errors="ignore")  # Read and decode file content
#         # st.write(f"File content preview (first 100 characters): {file_content[:100]}")
# else:
#     st.write("No files uploaded yet.")
# # if uploaded_file:
#     faq_text, links = process_faq_file(uploaded_file)
#     if "Unsupported file format" in faq_text:
#         st.error(faq_text)
#     else:
#         # st.success("Document uploaded and processed.")
#         st.session_state.faq_text = faq_text
#         st.session_state.links = links
#
#         # Display the uploaded file's name and a success message in the sidebar
#         st.sidebar.write(f"Uploaded file: {uploaded_file.name}")
#         st.sidebar.write(f"Number of paragraphs in FAQ: {len(faq_text.splitlines())}")
# Main Streamlit App
st.sidebar.title("Settings")
uploaded_files = st.sidebar.file_uploader("Upload multiple files", type=["docx", "pdf", "txt"], accept_multiple_files=True)
email = None
if "email" not in st.session_state:
    email = login()

if email or "email" in st.session_state:
    email = email or st.session_state.email

    # st.sidebar.title("Settings")
    # uploaded_files = st.sidebar.file_uploader("Upload multiple files", type=["docx", "pdf", "txt"], accept_multiple_files=True)

    if uploaded_files:
        combined_text = ""
        for uploaded_file in uploaded_files:
            if uploaded_file.name.endswith('.docx'):
                faq_text, links = process_faq_file(uploaded_file)
            elif uploaded_file.name.endswith('.pdf'):
                faq_text, links = extract_text_and_links_from_pdf(uploaded_file)
            elif uploaded_file.name.endswith('.txt'):
                faq_text = uploaded_file.read().decode('utf-8')

            else:
                st.warning("Unsupported file format. Please upload a DOCX or PDF.")
                faq_text, links = "", []

            st.session_state.faq_text += faq_text
            combined_text += faq_text + "\n"
            st.sidebar.write(f"File name: {uploaded_file.name}")
        st.session_state.faq_text = combined_text
        st.session_state.document_index = create_document_index(combined_text)

    if st.session_state.faq_text and not st.session_state.exit:
        if st.session_state.conversation:
            for idx, (q, a) in enumerate(st.session_state.conversation, start=1):
                st.markdown(f"**Q{idx}:** <span class='question'>{q}</span>", unsafe_allow_html=True)
                st.markdown(f"**A{idx}:** <span class='response'>{a}</span>", unsafe_allow_html=True)

        query = st.text_input("Ask your next question (type 'exit' to finish):", key=len(st.session_state.conversation))

        if query:
            if query.lower() == "exit":
                st.session_state.exit = True
                st.markdown("<span class='exit-text'>Thank you for using the L&D Chatbot. Have a great day!</span>",
                            unsafe_allow_html=True)
            elif is_query_relevant(query, st.session_state.document_index):
                with st.spinner("Fetching the answer..."):
                    answer = get_answer_from_cohere(query, st.session_state.faq_text)
                st.session_state.conversation.append((query, answer))
                store_query(email, query, answer)
                st.session_state.rerun_trigger += 1
                # Add a button to manually trigger rerun
                if st.button("Get Answer"):
                    # Removed the experimental_rerun and implemented a session state clear
                    st.session_state.clear()  # Clears the session state
            else:
                fallback_message = "I'm sorry, I can only answer questions related to the uploaded document."
                st.session_state.conversation.append((query, fallback_message))
    else:
        # st.write("Upload your FAQ document and ask a question.")
        st.markdown("#### <span class='wel-text'> We are happy to see you back again Click below to start a conversation  </span>",
            unsafe_allow_html=True)
else:
    st.write("Please login to use the chatbot.")
# Check if FAQ text is available and user is not exiting
# if "faq_text" in st.session_state and not st.session_state.exit:
#     if st.session_state.conversation:
#         # for idx, (q, a) in enumerate(st.session_state.conversation, start=1):
#         #     st.write(f"**Q{idx}:** {q}")
#         #     st.write(f"**A{idx}:** {a}")
#         for idx, (q, a) in enumerate(st.session_state.conversation, start=1):
#             st.markdown(f"**Q{idx}:** <span class='question'>{q}</span>", unsafe_allow_html=True)
#             st.markdown(f"**A{idx}:** <span class='response'>{a}</span>", unsafe_allow_html=True)

    # Button to trigger getting an answer
    # if st.button("Get Answer"):
    # query = st.text_input("Ask your next question (type 'exit' to finish):", key=len(st.session_state.conversation))
    #
    # if query:
    #     if query.lower() == "exit":
    #         st.session_state.exit = True
    #         st.markdown("<span class='exit-text'>Thank you for using the L&D Chatbot. Have a great day!</span>",
    #                     unsafe_allow_html=True)
    #         # st.write("##### Thank you for using the L&D Chatbot. Have a great day!")
    #     else:
    #         with st.spinner("Fetching the answer..."):
    #             answer = get_answer_from_cohere(query, st.session_state.faq_text)
    #         st.session_state.conversation.append((query, answer))

            # Trigger the rerun by updating the session state
            # st.session_state.rerun_trigger += 1
            #
            # # Add a button to manually trigger rerun
            # if st.button("Get Answer"):
            #     # Removed the experimental_rerun and implemented a session state clear
            #     st.session_state.clear()  # Clears the session state
            #
            # # Clear the text input field after getting the answer
            # # st.text_input("Ask your next question (type 'exit' to finish):", value="", key="question_input")

# else:
#     # st.write("Upload your FAQ document and ask a question.")
#     st.markdown("#### <span class='wel-text'> We are happy to see you back again Click below to start a conversation  </span>",unsafe_allow_html=True)
# Add a button to start a new conversation
if st.session_state.exit:
    if st.button("Start New Conversation"):
        # Reset session state to start fresh
        st.session_state.conversation = []
        st.session_state.exit = False
        st.session_state.rerun_trigger = 0
        st.session_state.clear()  # Clears the session state to start fresh
